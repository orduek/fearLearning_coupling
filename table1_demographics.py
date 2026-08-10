from __future__ import annotations

import math
from dataclasses import dataclass
from pathlib import Path
from typing import Mapping, Sequence

import pandas as pd
from scipy import stats


@dataclass(frozen=True)
class Table1Outputs:
    csv_path: Path
    md_path: Path
    html_path: Path
    docx_path: Path | None = None


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------

def _fmt_mean_sd(x: pd.Series) -> str:
    x = pd.to_numeric(x, errors="coerce").dropna()
    if len(x) == 0:
        return "—"
    mean = float(x.mean())
    sd = float(x.std(ddof=1)) if len(x) > 1 else float("nan")
    if math.isnan(sd):
        return f"{mean:.1f}"
    return f"{mean:.1f} ({sd:.1f})"


def _fmt_median_iqr(x: pd.Series) -> str:
    x = pd.to_numeric(x, errors="coerce").dropna()
    if len(x) == 0:
        return "—"
    q1 = float(x.quantile(0.25))
    med = float(x.quantile(0.50))
    q3 = float(x.quantile(0.75))
    return f"{med:.1f} ({q1:.1f}, {q3:.1f})"


def _fmt_n_pct(mask: pd.Series, denom: int | None = None) -> str:
    mask = mask.fillna(False).astype(bool)
    n = int(mask.sum())
    d = int(len(mask)) if denom is None else denom
    if d == 0:
        return "0 (0%)"
    return f"{n} ({(100.0 * n / d):.0f}%)"


def _fmt_p(p: float | None) -> str:
    if p is None or (isinstance(p, float) and math.isnan(p)):
        return ""
    if p < 0.001:
        return "<0.001"
    return f"{p:.3f}"


# ---------------------------------------------------------------------------
# Statistical tests (mirrors gtsummary::add_p() defaults)
# ---------------------------------------------------------------------------

def _kruskal_p(df: pd.DataFrame, var_col: str, group_col: str) -> float | None:
    groups = [
        pd.to_numeric(g[var_col], errors="coerce").dropna()
        for _, g in df.groupby(group_col, observed=True)
    ]
    groups = [g for g in groups if len(g) > 0]
    if len(groups) < 2:
        return None
    try:
        _, p = stats.kruskal(*groups)
        return float(p)
    except ValueError:
        return None


def _chi2_p(df: pd.DataFrame, var_col: str, group_col: str) -> float | None:
    table = pd.crosstab(df[var_col], df[group_col])
    if table.shape[0] < 2 or table.shape[1] < 2:
        return None
    try:
        _, p, _, _ = stats.chi2_contingency(table)
        return float(p)
    except ValueError:
        return None


# ---------------------------------------------------------------------------
# Table builder
# ---------------------------------------------------------------------------

def build_table1_demographics(
    csv_path: str | Path = "data/merged_caps_demographics.csv",
    *,
    id_col: str = "sub_id",
    group_col: str = "group",
    group_order: Sequence[str] = ("HC", "VCC", "VPTSD"),
    group_label_map: Mapping[str, str] | None = None,
    age_col: str = "Age",
    gender_col: str = "Gender",
    gender_map: Mapping[float, str] | None = None,
    caps_col: str = "total_caps",
    caps_label: str = "CAPS-5 Total",
    include_p_value: bool = True,
    output_dir: str | Path = "outputs",
    filename_stem: str = "table1_demographics",
) -> tuple[pd.DataFrame, Table1Outputs]:
    """Build a gtsummary-style Table 1 from `merged_caps_demographics.csv`.

    Columns: Overall + one per group (+ p-value). Continuous variables are
    reported as requested (Age: mean (SD); CAPS total: median (IQR), since it
    is right-skewed with a floor at 0). Categorical variables are n (%).
    p-values use Kruskal-Wallis (continuous) and chi-square (categorical),
    matching gtsummary::add_p() defaults for >2 groups.
    """
    csv_path = Path(csv_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    df = pd.read_csv(csv_path)

    if gender_map is None:
        gender_map = {1: "Male", 2: "Female", 1.0: "Male", 2.0: "Female"}
    if group_label_map is None:
        group_label_map = {"HC": "Healthy Controls", "VCC": "Combat Controls", "VPTSD": "PTSD"}

    df["_gender_label"] = df[gender_col].map(gender_map).fillna(df[gender_col].astype(str))

    observed = list(pd.Series(df[group_col].dropna().unique()).astype(str))
    ordered_groups = [g for g in group_order if g in observed] + [g for g in observed if g not in group_order]

    col_keys = ["Overall", *ordered_groups]
    col_labels = {
        "Overall": f"Overall, N = {df[id_col].nunique()}",
        **{
            g: f"{group_label_map.get(g, g)}, N = {int((df[group_col] == g).sum())}"
            for g in ordered_groups
        },
    }

    def subset_for(key: str) -> pd.DataFrame:
        return df if key == "Overall" else df[df[group_col] == key]

    # rows: list of (row_label, {col_key: value}, indent: bool)
    rows: list[tuple[str, dict[str, str], bool]] = []

    # --- Age ---
    age_row = {k: _fmt_mean_sd(subset_for(k)[age_col]) for k in col_keys}
    p_age = _fmt_p(_kruskal_p(df, age_col, group_col)) if include_p_value else None
    rows.append((f"{age_col}, Mean (SD)", age_row, False))

    # --- Sex ---
    sex_header = {k: "" for k in col_keys}
    rows.append(("Sex", sex_header, False))
    unique_labels = [x for x in pd.Series(df["_gender_label"].dropna().unique()).astype(str)]
    preferred = ["Male", "Female"]
    ordered_labels = [x for x in preferred if x in unique_labels] + sorted(
        [x for x in unique_labels if x not in preferred]
    )
    p_sex = _fmt_p(_chi2_p(df, "_gender_label", group_col)) if include_p_value else None
    for lab in ordered_labels:
        lab_row = {}
        for k in col_keys:
            sdf = subset_for(k)
            lab_row[k] = _fmt_n_pct(sdf["_gender_label"].astype(str) == lab, denom=len(sdf))
        rows.append((lab, lab_row, True))

    # --- CAPS total ---
    caps_row = {k: _fmt_median_iqr(subset_for(k)[caps_col]) for k in col_keys}
    p_caps = _fmt_p(_kruskal_p(df, caps_col, group_col)) if include_p_value else None
    rows.append((f"{caps_label}, Median (IQR)", caps_row, False))

    # --- assemble table ---
    display_cols = [col_labels[k] for k in col_keys]
    p_by_row_label = {
        f"{age_col}, Mean (SD)": p_age,
        "Sex": p_sex,
        f"{caps_label}, Median (IQR)": p_caps,
    }

    records = []
    index_labels = []
    indent_flags = []
    for label, values, indented in rows:
        display_label = f"    {label}" if indented else label
        index_labels.append(display_label)
        indent_flags.append(indented)
        rec = {display_cols[i]: values[k] for i, k in enumerate(col_keys)}
        if include_p_value:
            rec["p-value"] = p_by_row_label.get(label, "")
        records.append(rec)

    table = pd.DataFrame(records, index=index_labels)
    table.index.name = "Characteristic"

    csv_out = output_dir / f"{filename_stem}.csv"
    md_out = output_dir / f"{filename_stem}.md"
    html_out = output_dir / f"{filename_stem}.html"
    docx_out = output_dir / f"{filename_stem}.docx"

    table.to_csv(csv_out)
    md_out.write_text(table.to_markdown(), encoding="utf-8")
    html_out.write_text(_to_gtsummary_html(table, indent_flags, include_p_value), encoding="utf-8")

    wrote_docx = _write_docx(table, indent_flags, docx_out)

    return table, Table1Outputs(
        csv_path=csv_out,
        md_path=md_out,
        html_path=html_out,
        docx_path=(docx_out if wrote_docx else None),
    )


def _to_gtsummary_html(table: pd.DataFrame, indent_flags: list[bool], has_p: bool) -> str:
    style = """
    <style>
      table.t1 { border-collapse: collapse; font-family: -apple-system, Helvetica, Arial, sans-serif;
                 font-size: 14px; color: #1a1a1a; }
      table.t1 caption { font-weight: 600; text-align: left; padding-bottom: 6px; }
      table.t1 th, table.t1 td { padding: 5px 14px; text-align: left; white-space: nowrap; }
      table.t1 thead tr { border-top: 2px solid #333; border-bottom: 1px solid #333; }
      table.t1 thead th { font-weight: 600; }
      table.t1 tbody tr:last-child { border-bottom: 2px solid #333; }
      table.t1 td.stat, table.t1 th.stat { text-align: center; }
      table.t1 td.label-indent { padding-left: 32px; }
      table.t1 tfoot td { font-size: 11px; color: #555; padding-top: 8px; white-space: normal; }
    </style>
    """
    cols = list(table.columns)
    thead = "<tr><th></th>" + "".join(f'<th class="stat">{c}</th>' for c in cols) + "</tr>"
    body_rows = []
    for indented, (idx, row) in zip(indent_flags, table.iterrows()):
        label = idx.strip()
        label_cls = ' class="label-indent"' if indented else ""
        cells = "".join(f'<td class="stat">{"" if pd.isna(v) else v}</td>' for v in row[cols])
        body_rows.append(f"<tr><td{label_cls}>{label}</td>{cells}</tr>")
    footnote = ""
    if has_p:
        footnote = (
            "<tfoot><tr><td colspan='{}'>Statistics: Mean (SD); Median (IQR); n (%). "
            "p-values: Kruskal-Wallis rank sum test (continuous); Pearson's chi-squared test (categorical).</td></tr></tfoot>"
        ).format(len(cols) + 1)
    return (
        style
        + '<table class="t1"><caption>Table 1. Sample Characteristics</caption>'
        + f"<thead>{thead}</thead><tbody>{''.join(body_rows)}</tbody>{footnote}</table>"
    )


def _write_docx(table: pd.DataFrame, indent_flags: list[bool], docx_path: Path) -> bool:
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        return False

    doc = Document()
    doc.add_heading("Table 1. Sample Characteristics", level=1)

    t = doc.add_table(rows=1, cols=len(table.columns) + 1)
    t.style = "Table Grid"

    hdr = t.rows[0].cells
    hdr[0].text = table.index.name or "Characteristic"
    for j, col in enumerate(table.columns, start=1):
        hdr[j].text = str(col)

    for indented, (idx, row) in zip(indent_flags, table.iterrows()):
        cells = t.add_row().cells
        para = cells[0].paragraphs[0]
        para.text = str(idx).strip()
        if indented:
            para.paragraph_format.left_indent = Inches(0.25)
        for j, col in enumerate(table.columns, start=1):
            val = row[col]
            cells[j].text = "" if pd.isna(val) else str(val)

    doc.save(docx_path)
    return True


if __name__ == "__main__":
    table, paths = build_table1_demographics()
    print(table)
    print("\nWrote:")
    print(f"- {paths.csv_path}")
    print(f"- {paths.md_path}")
    print(f"- {paths.html_path}")
    if paths.docx_path is not None:
        print(f"- {paths.docx_path}")
    else:
        print("- (DOCX not written; install `python-docx` to enable Word export)")
